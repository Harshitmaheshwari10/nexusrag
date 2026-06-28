"""
NexusRAG Document Processor
Supports: PDF, TXT, MD, DOCX, CSV
Strategy: Semantic chunking (sentence-boundary aware, not fixed-size)
"""

import re
import io
from typing import List, Tuple, Dict, Any
from dataclasses import dataclass


@dataclass
class Chunk:
    text: str
    metadata: Dict[str, Any]


class DocumentProcessor:
    """
    Processes documents into semantically-aware chunks.
    Uses sentence-boundary chunking with overlap.
    """

    def __init__(self, chunk_size: int = 400, overlap: int = 80):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def process(self, uploaded_file) -> Tuple[List[str], Dict[str, Any]]:
        """Extract text and return (chunks, metadata)."""
        file_type = uploaded_file.type or ""
        name = uploaded_file.name
        raw_bytes = uploaded_file.read()

        if "pdf" in file_type or name.endswith(".pdf"):
            text = self._extract_pdf(raw_bytes)
        elif name.endswith(".docx"):
            text = self._extract_docx(raw_bytes)
        elif name.endswith(".csv"):
            text = self._extract_csv(raw_bytes)
        else:
            # TXT / MD / fallback
            text = raw_bytes.decode("utf-8", errors="replace")

        chunks = self._semantic_chunk(text)
        metadata = {
            "filename": name,
            "file_type": file_type,
            "total_chars": len(text),
            "num_chunks": len(chunks),
        }
        return chunks, metadata

    # ── Extractors ────────────────────────────────────────────────────────────

    def _extract_pdf(self, raw_bytes: bytes) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"[Page {i+1}]\n{page_text}")
            return "\n\n".join(pages)
        except ImportError:
            # Fallback: try pdfminer
            try:
                from pdfminer.high_level import extract_text_to_fp
                from pdfminer.layout import LAParams
                output = io.StringIO()
                extract_text_to_fp(io.BytesIO(raw_bytes), output, laparams=LAParams())
                return output.getvalue()
            except Exception:
                return "[PDF extraction failed — install pypdf or pdfminer.six]"
        except Exception as e:
            return f"[PDF error: {e}]"

    def _extract_docx(self, raw_bytes: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except Exception as e:
            return f"[DOCX error: {e}]"

    def _extract_csv(self, raw_bytes: bytes) -> str:
        try:
            import csv
            text = raw_bytes.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
            if not rows:
                return ""
            header = rows[0]
            lines = [", ".join(header)]
            for row in rows[1:]:
                if any(cell.strip() for cell in row):
                    lines.append(", ".join(f"{h}: {v}" for h, v in zip(header, row)))
            return "\n".join(lines)
        except Exception as e:
            return f"[CSV error: {e}]"

    # ── Semantic Chunker ──────────────────────────────────────────────────────

    def _semantic_chunk(self, text: str) -> List[str]:
        """
        Split text into semantically coherent chunks:
        1. Split on sentence boundaries
        2. Group sentences until chunk_size reached
        3. Add overlap from previous chunk
        """
        # Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)

        # Split into sentences
        sentences = self._split_sentences(text)
        if not sentences:
            return [text[:self.chunk_size]] if text else []

        chunks = []
        current_chunk = []
        current_len = 0
        overlap_buffer = []

        for sent in sentences:
            sent_len = len(sent)

            # If a single sentence is longer than chunk_size, hard split it
            if sent_len > self.chunk_size:
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    overlap_buffer = current_chunk[-2:]
                    current_chunk = []
                    current_len = 0
                # Hard split the long sentence
                words = sent.split()
                temp = []
                temp_len = 0
                for w in words:
                    if temp_len + len(w) + 1 > self.chunk_size and temp:
                        chunks.append(" ".join(temp))
                        overlap_buffer = temp[-10:]
                        temp = overlap_buffer[:]
                        temp_len = sum(len(x)+1 for x in temp)
                    temp.append(w)
                    temp_len += len(w) + 1
                if temp:
                    current_chunk = temp
                    current_len = temp_len
                continue

            if current_len + sent_len + 1 > self.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunks.append(chunk_text)
                # Build overlap
                overlap_text = []
                overlap_len = 0
                for s in reversed(current_chunk):
                    if overlap_len + len(s) > self.overlap:
                        break
                    overlap_text.insert(0, s)
                    overlap_len += len(s) + 1
                current_chunk = overlap_text
                current_len = overlap_len

            current_chunk.append(sent)
            current_len += sent_len + 1

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        # Filter empty / too-short
        chunks = [c.strip() for c in chunks if len(c.strip()) > 30]
        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        """Split on sentence boundaries, preserving paragraph breaks."""
        # Split on paragraphs first
        paragraphs = re.split(r'\n\n+', text)
        sentences = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            # Split on sentence endings
            sents = re.split(r'(?<=[.!?])\s+(?=[A-Z\[\(])', para)
            sentences.extend([s.strip() for s in sents if s.strip()])
        return sentences
